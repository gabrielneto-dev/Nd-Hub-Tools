#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════╗
║          CONVERTIX - Universal File Converter                ║
║          Clone profissional do Convertio.co                  ║
╚══════════════════════════════════════════════════════════════╝
"""

import os
import sys
import subprocess
import shutil
import json
from pathlib import Path
from typing import Optional

# ─── Cores ANSI ────────────────────────────────────────────────
R  = "\033[0m"
B  = "\033[1m"
DIM = "\033[2m"
CY = "\033[96m"
GR = "\033[92m"
YE = "\033[93m"
RE = "\033[91m"
MA = "\033[95m"
BL = "\033[94m"
WH = "\033[97m"
BG_BL = "\033[44m"
BG_CY = "\033[46m"
BG_GR = "\033[42m"
BG_RE = "\033[41m"

# ─── Mapa de conversões suportadas ─────────────────────────────
CONVERSION_MAP = {
    # ── IMAGENS ──────────────────────────────────────────────
    "image": {
        "input":  [".jpg",".jpeg",".png",".gif",".bmp",".tiff",".tif",
                   ".webp",".ico",".svg",".heic",".avif",".eps",".psd",
                   ".cr2",".nef",".dng",".raw"],
        "output": [".jpg",".jpeg",".png",".gif",".bmp",".tiff",".tif",
                   ".webp",".ico",".pdf",".eps"],
        "engine": "pillow+imagemagick"
    },
    # ── VÍDEO ────────────────────────────────────────────────
    "video": {
        "input":  [".mp4",".avi",".mov",".mkv",".wmv",".flv",".webm",
                   ".m4v",".3gp",".ogv",".ts",".mts",".m2ts",".vob",".f4v"],
        "output": [".mp4",".avi",".mov",".mkv",".wmv",".flv",".webm",
                   ".m4v",".3gp",".ogv",".gif",".mp3",".aac",".wav"],
        "engine": "ffmpeg"
    },
    # ── ÁUDIO ────────────────────────────────────────────────
    "audio": {
        "input":  [".mp3",".wav",".flac",".aac",".ogg",".m4a",".wma",
                   ".opus",".aiff",".au",".mid",".midi",".amr"],
        "output": [".mp3",".wav",".flac",".aac",".ogg",".m4a",".opus",".aiff"],
        "engine": "ffmpeg"
    },
    # ── DOCUMENTOS ───────────────────────────────────────────
    "document": {
        "input":  [".pdf",".docx",".doc",".txt",".rtf",".odt",
                   ".html",".htm",".md",".csv",".xlsx",".xls",".json",".xml"],
        "output": [".pdf",".docx",".txt",".html",".md",".csv",
                   ".xlsx",".json",".xml"],
        "engine": "mixed"
    },
    # ── DADOS ────────────────────────────────────────────────
    "data": {
        "input":  [".csv",".xlsx",".xls",".json",".xml",".tsv",".yaml",".yml"],
        "output": [".csv",".xlsx",".json",".xml",".tsv",".html",".md"],
        "engine": "pandas"
    },
}

# Mapa inverso: extensão → categoria
EXT_TO_CAT = {}
for cat, info in CONVERSION_MAP.items():
    for ext in info["input"]:
        EXT_TO_CAT[ext] = cat


def clear():
    os.system("cls" if os.name == "nt" else "clear")


def banner():
    print(f"""
{CY}{B}
  ██████╗ ██████╗ ███╗   ██╗██╗   ██╗███████╗██████╗ ████████╗██╗██╗  ██╗
 ██╔════╝██╔═══██╗████╗  ██║██║   ██║██╔════╝██╔══██╗╚══██╔══╝██║╚██╗██╔╝
 ██║     ██║   ██║██╔██╗ ██║██║   ██║█████╗  ██████╔╝   ██║   ██║ ╚███╔╝
 ██║     ██║   ██║██║╚██╗██║╚██╗ ██╔╝██╔══╝  ██╔══██╗   ██║   ██║ ██╔██╗
 ╚██████╗╚██████╔╝██║ ╚████║ ╚████╔╝ ███████╗██║  ██║   ██║   ██║██╔╝ ██╗
  ╚═════╝ ╚═════╝ ╚═╝  ╚═══╝  ╚═══╝  ╚══════╝╚═╝  ╚═╝   ╚═╝   ╚═╝╚═╝  ╚═╝
{R}{DIM}  Conversor Universal de Arquivos  ·  v2.0  ·  Powered by Python{R}
{CY}{'─'*74}{R}""")


def info(msg):  print(f"  {BL}ℹ{R}  {msg}")
def ok(msg):    print(f"  {GR}✔{R}  {msg}")
def warn(msg):  print(f"  {YE}⚠{R}  {WH}{msg}{R}")
def err(msg):   print(f"  {RE}✘{R}  {RE}{msg}{R}")
def step(msg):  print(f"\n  {MA}▶{R}  {B}{msg}{R}")
def sep():      print(f"  {DIM}{'─'*60}{R}")


def progress_bar(current, total, width=40):
    pct = current / total if total else 1
    filled = int(width * pct)
    bar = f"{GR}{'█' * filled}{DIM}{'░' * (width - filled)}{R}"
    return f"  [{bar}] {CY}{int(pct*100):3d}%{R}"


def get_category(ext: str) -> Optional[str]:
    return EXT_TO_CAT.get(ext.lower())


def get_available_outputs(input_ext: str) -> list:
    cat = get_category(input_ext)
    if not cat:
        return []
    return [e for e in CONVERSION_MAP[cat]["output"] if e != input_ext.lower()]


# ─── ENGINES DE CONVERSÃO ──────────────────────────────────────

def convert_image(src: Path, dst: Path) -> bool:
    """Converte imagens usando Pillow + ImageMagick como fallback."""
    try:
        from PIL import Image
        src_ext = src.suffix.lower()
        dst_ext = dst.suffix.lower()

        # Casos especiais: SVG, EPS, PSD → ImageMagick
        if src_ext in [".svg", ".eps", ".psd", ".heic", ".cr2", ".nef", ".dng"] or \
           dst_ext in [".svg", ".eps"]:
            return _imagemagick(src, dst)

        img = Image.open(src)

        # Tratar transparência
        if dst_ext in [".jpg", ".jpeg"]:
            if img.mode in ("RGBA", "LA", "P"):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                if img.mode in ("RGBA", "LA"):
                    bg.paste(img, mask=img.split()[-1])
                else:
                    bg.paste(img)
                img = bg
            else:
                img = img.convert("RGB")

        # PDF especial
        if dst_ext == ".pdf":
            if img.mode in ("RGBA", "LA", "P"):
                img = img.convert("RGB")
            img.save(dst, "PDF", resolution=150)
        elif dst_ext in [".tiff", ".tif"]:
            img.save(dst, "TIFF")
        elif dst_ext == ".webp":
            img.save(dst, "WEBP", quality=85)
        elif dst_ext == ".ico":
            img.save(dst, "ICO", sizes=[(256,256),(128,128),(64,64),(32,32),(16,16)])
        else:
            img.save(dst)
        return True
    except Exception as e:
        warn(f"Pillow falhou ({e}), tentando ImageMagick...")
        return _imagemagick(src, dst)


def _imagemagick(src: Path, dst: Path) -> bool:
    result = subprocess.run(
        ["convert", str(src), str(dst)],
        capture_output=True, text=True
    )
    return result.returncode == 0


def convert_media(src: Path, dst: Path, media_type: str) -> bool:
    """Converte áudio/vídeo via ffmpeg."""
    dst_ext = dst.suffix.lower()
    cmd = ["ffmpeg", "-i", str(src), "-y"]

    if media_type == "video":
        if dst_ext in [".mp4", ".m4v"]:
            cmd += ["-c:v", "libx264", "-crf", "23", "-c:a", "aac", "-b:a", "192k"]
        elif dst_ext == ".webm":
            cmd += ["-c:v", "libvpx-vp9", "-b:v", "0", "-crf", "30", "-c:a", "libopus"]
        elif dst_ext == ".gif":
            cmd = ["ffmpeg", "-i", str(src), "-vf",
                   "fps=10,scale=640:-1:flags=lanczos,split[s0][s1];[s0]palettegen[p];[s1][p]paletteuse",
                   "-y", str(dst)]
            result = subprocess.run(cmd, capture_output=True, text=True)
            return result.returncode == 0
        elif dst_ext in [".mp3", ".aac", ".wav"]:
            cmd += ["-vn"]  # só áudio
        elif dst_ext in [".avi"]:
            cmd += ["-c:v", "mpeg4", "-c:a", "mp3"]
        elif dst_ext == ".mkv":
            cmd += ["-c:v", "copy", "-c:a", "copy"]
        else:
            cmd += ["-c:v", "libx264", "-c:a", "aac"]

    elif media_type == "audio":
        if dst_ext == ".mp3":
            cmd += ["-c:a", "libmp3lame", "-b:a", "192k"]
        elif dst_ext == ".flac":
            cmd += ["-c:a", "flac"]
        elif dst_ext == ".wav":
            cmd += ["-c:a", "pcm_s16le"]
        elif dst_ext == ".ogg":
            cmd += ["-c:a", "libvorbis", "-q:a", "4"]
        elif dst_ext == ".aac":
            cmd += ["-c:a", "aac", "-b:a", "192k"]
        elif dst_ext == ".opus":
            cmd += ["-c:a", "libopus", "-b:a", "128k"]
        elif dst_ext in [".m4a"]:
            cmd += ["-c:a", "aac", "-b:a", "192k"]

    cmd.append(str(dst))
    result = subprocess.run(cmd, capture_output=True, text=True)
    return result.returncode == 0


def convert_document(src: Path, dst: Path) -> bool:
    """Converte documentos."""
    src_ext = src.suffix.lower()
    dst_ext = dst.suffix.lower()

    # ── TXT → vários ──
    if src_ext == ".txt":
        text = src.read_text(encoding="utf-8", errors="replace")
        if dst_ext == ".md":
            dst.write_text(text, encoding="utf-8"); return True
        if dst_ext == ".html":
            html = f"<!DOCTYPE html><html><body><pre>{text}</pre></body></html>"
            dst.write_text(html, encoding="utf-8"); return True
        if dst_ext == ".pdf":
            return _text_to_pdf(text, dst)
        if dst_ext == ".docx":
            return _text_to_docx(text, dst)
        if dst_ext == ".json":
            import json
            dst.write_text(json.dumps({"content": text}, indent=2), encoding="utf-8"); return True

    # ── MD → vários ──
    if src_ext == ".md":
        text = src.read_text(encoding="utf-8", errors="replace")
        if dst_ext == ".html":
            html = _md_to_html(text)
            dst.write_text(html, encoding="utf-8"); return True
        if dst_ext == ".txt":
            import re
            plain = re.sub(r'[#*`_\[\]()!>~\-]+', '', text)
            dst.write_text(plain, encoding="utf-8"); return True
        if dst_ext == ".pdf":
            html = _md_to_html(text)
            tmp = src.with_suffix(".html.tmp")
            tmp.write_text(html, encoding="utf-8")
            ok_r = _html_to_pdf(tmp, dst)
            tmp.unlink(missing_ok=True)
            return ok_r
        if dst_ext == ".docx":
            return _text_to_docx(text, dst)

    # ── HTML → vários ──
    if src_ext in [".html", ".htm"]:
        text = src.read_text(encoding="utf-8", errors="replace")
        if dst_ext == ".txt":
            import re
            plain = re.sub(r'<[^>]+>', '', text)
            dst.write_text(plain, encoding="utf-8"); return True
        if dst_ext == ".md":
            import re
            plain = re.sub(r'<[^>]+>', '', text)
            dst.write_text(plain, encoding="utf-8"); return True
        if dst_ext == ".pdf":
            return _html_to_pdf(src, dst)

    # ── PDF → TXT ──
    if src_ext == ".pdf" and dst_ext == ".txt":
        return _pdf_to_txt(src, dst)

    # ── DOCX → TXT / PDF ──
    if src_ext == ".docx":
        if dst_ext == ".txt":
            return _docx_to_txt(src, dst)
        if dst_ext == ".pdf":
            return _docx_to_pdf(src, dst)
        if dst_ext == ".html":
            return _docx_to_html(src, dst)

    # ── JSON ↔ XML ──  (usa helpers dedicados, não pandas)
    if src_ext == ".json" and dst_ext == ".xml":
        ok_r = _json_to_xml(src, dst)
        return ok_r
    if src_ext == ".xml" and dst_ext == ".json":
        ok_r = _xml_to_json(src, dst)
        return ok_r

    # ── JSON → HTML ──
    if src_ext == ".json" and dst_ext == ".html":
        data = json.loads(src.read_text())
        html = f"<html><body><pre>{json.dumps(data, indent=2, ensure_ascii=False)}</pre></body></html>"
        dst.write_text(html); return True

    # ── QUALQUER → TXT fallback ──
    if dst_ext == ".txt":
        try:
            dst.write_text(src.read_text(encoding="utf-8", errors="replace"))
            return True
        except: pass

    warn(f"Conversão {src_ext} → {dst_ext} não suportada diretamente.")
    return False


def convert_data(src: Path, dst: Path) -> bool:
    """Converte arquivos de dados via pandas."""
    import pandas as pd
    src_ext = src.suffix.lower()
    dst_ext = dst.suffix.lower()

    try:
        # ── Leitura ──
        if src_ext == ".csv":
            df = pd.read_csv(src)
        elif src_ext in [".xlsx", ".xls"]:
            df = pd.read_excel(src)
        elif src_ext == ".json":
            df = pd.read_json(src)
        elif src_ext in [".tsv"]:
            df = pd.read_csv(src, sep="\t")
        elif src_ext == ".xml":
            df = pd.read_xml(src)
        elif src_ext in [".yaml", ".yml"]:
            import yaml
            with open(src) as f:
                data = yaml.safe_load(f)
            df = pd.json_normalize(data) if isinstance(data, list) else pd.DataFrame([data])
        else:
            df = pd.read_csv(src)

        # ── Escrita ──
        if dst_ext == ".csv":
            df.to_csv(dst, index=False)
        elif dst_ext == ".xlsx":
            df.to_excel(dst, index=False, engine="openpyxl")
        elif dst_ext == ".json":
            df.to_json(dst, orient="records", indent=2, force_ascii=False)
        elif dst_ext == ".tsv":
            df.to_csv(dst, sep="\t", index=False)
        elif dst_ext == ".html":
            df.to_html(dst, index=False, border=1)
        elif dst_ext == ".xml":
            df.to_xml(dst, index=False)
        elif dst_ext == ".md":
            dst.write_text(df.to_markdown(index=False))
        else:
            df.to_csv(dst, index=False)
        return True

    except Exception as e:
        err(f"Erro pandas: {e}")
        return False


# ── Helpers de documentos ──────────────────────────────────────

def _text_to_pdf(text: str, dst: Path) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(dst), pagesize=A4)
        w, h = A4
        y = h - 50
        for line in text.split("\n"):
            if y < 50:
                c.showPage(); y = h - 50
            c.drawString(40, y, line[:110])
            y -= 14
        c.save()
        return True
    except:
        pass
    # Fallback: ImageMagick
    try:
        result = subprocess.run(
            ["convert", "-size", "2480x3507", "xc:white",
             "-font", "DejaVu-Sans", "-pointsize", "14",
             "-draw", f"text 40,60 '{text[:200]}'",
             str(dst)], capture_output=True)
        return result.returncode == 0
    except:
        return False


def _text_to_docx(text: str, dst: Path) -> bool:
    try:
        from docx import Document
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(dst))
        return True
    except Exception as e:
        err(f"python-docx: {e}"); return False


def _md_to_html(text: str) -> str:
    """Converte Markdown básico para HTML."""
    import re
    html = text
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.M)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.M)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.M)
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
    html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)
    html = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', html)
    lines = html.split("\n")
    result = []
    for l in lines:
        if l.startswith("<h") or l.startswith("<"):
            result.append(l)
        elif l.strip():
            result.append(f"<p>{l}</p>")
        else:
            result.append("<br>")
    return "<!DOCTYPE html><html><head><meta charset='utf-8'></head><body>\n" + \
           "\n".join(result) + "\n</body></html>"


def _html_to_pdf(src: Path, dst: Path) -> bool:
    for tool in ["wkhtmltopdf", "chromium-browser", "google-chrome", "chromium"]:
        path = shutil.which(tool)
        if path:
            if "wkhtmltopdf" in tool:
                r = subprocess.run([path, str(src), str(dst)], capture_output=True)
            else:
                r = subprocess.run([path, "--headless", "--print-to-pdf="+str(dst),
                                    str(src)], capture_output=True)
            if r.returncode == 0:
                return True
    # Fallback: ImageMagick via screenshot (muito básico)
    return _imagemagick(src, dst)


def _pdf_to_txt(src: Path, dst: Path) -> bool:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(src))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        dst.write_text(text, encoding="utf-8")
        return True
    except Exception as e:
        err(f"pypdf: {e}"); return False


def _docx_to_txt(src: Path, dst: Path) -> bool:
    try:
        from docx import Document
        doc = Document(str(src))
        text = "\n".join(p.text for p in doc.paragraphs)
        dst.write_text(text, encoding="utf-8")
        return True
    except Exception as e:
        err(f"python-docx: {e}"); return False


def _docx_to_pdf(src: Path, dst: Path) -> bool:
    for tool in ["libreoffice", "soffice"]:
        path = shutil.which(tool)
        if path:
            r = subprocess.run([path, "--headless", "--convert-to", "pdf",
                                "--outdir", str(dst.parent), str(src)],
                               capture_output=True)
            if r.returncode == 0:
                return True
    warn("LibreOffice não encontrado. Convertendo texto do docx para PDF...")
    ok_ = _docx_to_txt(src, src.with_suffix(".txt.tmp"))
    if ok_:
        tmp = src.with_suffix(".txt.tmp")
        text = tmp.read_text(encoding="utf-8", errors="replace")
        r = _text_to_pdf(text, dst)
        tmp.unlink(missing_ok=True)
        return r
    return False


def _docx_to_html(src: Path, dst: Path) -> bool:
    try:
        from docx import Document
        doc = Document(str(src))
        paragraphs = [f"<p>{p.text}</p>" for p in doc.paragraphs if p.text.strip()]
        html = "<!DOCTYPE html><html><body>\n" + "\n".join(paragraphs) + "\n</body></html>"
        dst.write_text(html, encoding="utf-8")
        return True
    except Exception as e:
        err(f"docx→html: {e}"); return False


def _json_to_xml(src: Path, dst: Path) -> bool:
    import xml.etree.ElementTree as ET
    def dict_to_xml(d, parent):
        if isinstance(d, dict):
            for k, v in d.items():
                child = ET.SubElement(parent, str(k).replace(" ","_"))
                dict_to_xml(v, child)
        elif isinstance(d, list):
            for item in d:
                child = ET.SubElement(parent, "item")
                dict_to_xml(item, child)
        else:
            parent.text = str(d)
    try:
        data = json.loads(src.read_text())
        root = ET.Element("root")
        dict_to_xml(data, root)
        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ")
        tree.write(str(dst), encoding="unicode", xml_declaration=True)
        return True
    except Exception as e:
        err(f"json→xml: {e}"); return False


def _xml_to_json(src: Path, dst: Path) -> bool:
    import xml.etree.ElementTree as ET
    def xml_to_dict(elem):
        d = {}
        if elem.attrib:
            d["@attributes"] = dict(elem.attrib)
        children = list(elem)
        if children:
            child_dict = {}
            for child in children:
                k = child.tag
                v = xml_to_dict(child)
                if k in child_dict:
                    if not isinstance(child_dict[k], list):
                        child_dict[k] = [child_dict[k]]
                    child_dict[k].append(v)
                else:
                    child_dict[k] = v
            d.update(child_dict)
        elif elem.text and elem.text.strip():
            d["text"] = elem.text.strip()
        return d
    try:
        tree = ET.parse(str(src))
        root = tree.getroot()
        data = {root.tag: xml_to_dict(root)}
        dst.write_text(json.dumps(data, indent=2, ensure_ascii=False))
        return True
    except Exception as e:
        err(f"xml→json: {e}"); return False


# ─── DISPATCHER PRINCIPAL ──────────────────────────────────────

def do_convert(src: Path, dst_ext: str, output_dir: Path) -> tuple[bool, Path]:
    """Converte src para dst_ext, salva em output_dir."""
    src_ext = src.suffix.lower()
    dst_name = src.stem + dst_ext
    dst = output_dir / dst_name

    cat = get_category(src_ext)

    if not cat:
        # Tenta ffmpeg como curinga
        cmd = ["ffmpeg", "-i", str(src), "-y", str(dst)]
        r = subprocess.run(cmd, capture_output=True)
        return r.returncode == 0, dst

    success = False
    if cat == "image":
        success = convert_image(src, dst)
    elif cat == "video":
        success = convert_media(src, dst, "video")
    elif cat == "audio":
        success = convert_media(src, dst, "audio")
    elif cat == "document":
        # Tenta pipeline de dados primeiro para formatos tabulares
        if src_ext in CONVERSION_MAP["data"]["input"] and \
           dst_ext in CONVERSION_MAP["data"]["output"]:
            success = convert_data(src, dst)
        if not success:
            success = convert_document(src, dst)
    elif cat == "data":
        success = convert_data(src, dst)

    return success, dst


# ─── INTERFACE CONSOLE ─────────────────────────────────────────

def _render_browser(cwd: Path, entries: list, selected: set[Path], filter_ext: str):
    """Renderiza o painel do navegador de arquivos."""
    clear()

    # ── Cabeçalho ──────────────────────────────────────────────
    print(f"\n{CY}{B}  ⬡ CONVERTIX  {R}{DIM}— Navegador de Arquivos{R}")
    print(f"  {DIM}{'─'*62}{R}")

    # Breadcrumb visual
    parts = list(cwd.parts)
    crumbs = []
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            crumbs.append(f"{CY}{B} {part} {R}")
        else:
            crumbs.append(f"{DIM}{part}{R}")
    print("  " + f"{DIM} ›{R} ".join(crumbs))

    # Stats da pasta
    dirs_  = [e for e in entries if Path(e).is_dir()]
    files_ = [e for e in entries if Path(e).is_file()]
    n_supp = sum(1 for e in files_ if get_category(Path(e).suffix.lower()))
    flt_txt = f"  {YE}● filtro: *{filter_ext}{R}" if filter_ext else ""
    print(f"  {DIM}  {len(dirs_)} pasta(s)  ·  {len(files_)} arquivo(s)  ·  {n_supp} suportado(s){R}{flt_txt}")
    print(f"  {DIM}{'─'*62}{R}")

    # Legenda
    cmds = [("[n]","selec/entrar"), ("[0]","subir"), ("[a]","todos"),
            ("[f]","filtro"), ("[x]","limpar"), ("[ok]","confirmar")]
    row = "  ".join(f"{DIM}[{c}]{R} {WH}{d}{R}" for c,d in [
        ("n","selec/entrar"),("0","subir"),("a","todos"),
        ("f","filtro"),("x","limpar"),("ok","confirmar")])
    print(f"  {DIM}  {row}{R}\n")

    # ── PASTAS ─────────────────────────────────────────────────
    i_global = 1
    if dirs_:
        print(f"  {BL}  📁  PASTAS{R}")
        for entry in dirs_:
            p = Path(entry)
            try:
                n_items = len(list(p.iterdir()))
                sub_info = f"  {DIM}{n_items} itens{R}"
            except PermissionError:
                sub_info = f"  {RE}sem acesso{R}"
            num = f"{CY}{i_global:>3}.{R}"
            print(f"{num}      {BL}{B}{p.name}/{R}{sub_info}")
            i_global += 1
        print()

    # ── ARQUIVOS ───────────────────────────────────────────────
    cat_badge = {
        "image":    (f"{MA}", "IMG"),
        "video":    (f"{RE}", "VID"),
        "audio":    (f"{GR}", "AUD"),
        "document": (f"{YE}", "DOC"),
        "data":     (f"{CY}", "DAT"),
    }
    if files_:
        print(f"  {WH}  📄  ARQUIVOS{R}")
        for entry in files_:
            p   = Path(entry)
            ext = p.suffix.lower()
            cat = get_category(ext)
            sz  = p.stat().st_size if p.exists() else 0
            num = f"{CY}{i_global:>3}.{R}"

            if p in selected:
                mark  = f"{GR}{B} ✔ {R}"
                fname = f"{GR}{B}{p.name}{R}"
            else:
                mark  = f"   "
                fname = f"{WH}{p.name}{R}" if cat else f"{DIM}{p.name}{R}"

            if cat:
                col, badge = cat_badge.get(cat, (WH, "???"))
                tag  = f"{col}{B}[{badge}]{R}"
                meta = f"{DIM} {ext}  {_human(sz)}{R}"
            else:
                tag  = f"{DIM}[ -- ]{R}"
                meta = f"{RE}{DIM} não suportado{R}"

            print(f"{num}{mark}{tag}  {fname}{meta}")
            i_global += 1
        print()

    if i_global == 1:
        print(f"\n  {DIM}  (pasta vazia){R}\n")

    # ── PAINEL DE SELECIONADOS ─────────────────────────────────
    print(f"  {DIM}{'─'*62}{R}")
    if selected:
        total_sz = sum(s.stat().st_size for s in selected if s.exists())
        print(f"  {GR}{B}  ✔ SELECIONADOS  ({len(selected)} arquivo(s)  ·  {_human(total_sz)}){R}")
        for idx, s in enumerate(sorted(selected, key=lambda x: x.name), 1):
            sz = s.stat().st_size if s.exists() else 0
            try:
                rel = s.relative_to(cwd)
                loc = f"{DIM}[aqui]{R}"
            except ValueError:
                loc = f"{DIM}[{s.parent.name}/]{R}"
            ext_badge = s.suffix.lower()
            print(f"    {GR}{idx}.{R}  {WH}{B}{s.name}{R}  {loc}  {DIM}{ext_badge}  {_human(sz)}{R}")
    else:
        print(f"  {DIM}  ○ Nenhum arquivo selecionado  —  navegue e pressione o número do arquivo{R}")
    print()



def browse_files() -> list[Path]:
    """Navegador interativo de arquivos a partir do diretório de execução."""
    cwd      = Path.cwd()
    selected : set[Path] = set()
    filter_ext = ""
    # força flush antes do clear() para o terminal não perder output
    sys.stdout.flush()

    while True:
        # Monta lista de entradas: primeiro ".." se não for raiz, depois dirs, depois arquivos
        raw_entries: list[Path] = []
        try:
            dirs  = sorted([p for p in cwd.iterdir() if p.is_dir()
                            and not p.name.startswith(".")],
                           key=lambda p: p.name.lower())
            files = sorted([p for p in cwd.iterdir() if p.is_file()
                            and not p.name.startswith(".")
                            and (not filter_ext or p.suffix.lower() == filter_ext)],
                           key=lambda p: p.name.lower())
            raw_entries = dirs + files
        except PermissionError:
            err("Sem permissão para acessar esta pasta.")
            cwd = cwd.parent
            continue

        # Adiciona entrada ".." para subir
        has_parent = cwd.parent != cwd
        display_entries = raw_entries  # índices 1-based

        _render_browser(cwd, display_entries, selected, filter_ext)

        cmd = input(f"  {CY}►{R} ").strip().lower()

        # ── Confirmar seleção ──
        if cmd in ("ok", "done", "confirmar", "c", ""):
            if selected:
                return list(selected)
            warn("Selecione ao menos 1 arquivo antes de confirmar.")
            input(f"  {DIM}Enter para continuar...{R}")
            continue

        # ── Subir um nível ──
        if cmd == "0":
            if has_parent:
                cwd = cwd.parent
            else:
                warn("Já está na raiz.")
            continue

        # ── Selecionar todos os arquivos visíveis ──
        if cmd == "a":
            for p in display_entries:
                if p.is_file():
                    cat = get_category(p.suffix.lower())
                    if cat:
                        selected.add(p)
            ok(f"{len(selected)} arquivo(s) selecionado(s).")
            input(f"  {DIM}Enter para continuar...{R}")
            continue

        # ── Limpar seleção ──
        if cmd == "x":
            selected.clear()
            continue

        # ── Filtrar por extensão ──
        if cmd == "f":
            raw = input(f"  {CY}Extensão para filtrar (ex: .jpg, ou vazio para limpar):{R} ").strip()
            filter_ext = raw if raw.startswith(".") else ("." + raw if raw else "")
            continue

        # ── Digitar caminho direto ──
        if cmd.startswith("/") or cmd.startswith("~") or (len(cmd) > 1 and cmd[1] == ":"):
            p = Path(os.path.expanduser(cmd))
            if p.is_dir():
                cwd = p
            elif p.is_file():
                selected.add(p)
                ok(f"Adicionado: {p.name}")
                input(f"  {DIM}Enter para continuar...{R}")
            else:
                err(f"Caminho não encontrado: {cmd}")
                input(f"  {DIM}Enter para continuar...{R}")
            continue

        # ── Selecionar/navegar por número ──
        # Suporta múltiplos: "1 3 5" ou "1,3,5" ou "1-5"
        tokens = cmd.replace(",", " ").split()
        indices: list[int] = []
        for tok in tokens:
            if "-" in tok and len(tok) > 1:
                try:
                    a, b = tok.split("-", 1)
                    indices += list(range(int(a), int(b)+1))
                except ValueError:
                    pass
            elif tok.isdigit():
                indices.append(int(tok))

        if not indices:
            warn("Comando não reconhecido. Digite um número, 'ok', '0', 'a', 'f' ou um caminho.")
            input(f"  {DIM}Enter para continuar...{R}")
            continue

        for idx in indices:
            if 1 <= idx <= len(display_entries):
                p = display_entries[idx - 1]
                if p.is_dir():
                    if len(indices) == 1:
                        cwd = p  # navega para a pasta
                    else:
                        warn(f"'{p.name}' é uma pasta — ignorada na seleção múltipla.")
                else:
                    cat = get_category(p.suffix.lower())
                    if p in selected:
                        selected.discard(p)
                    elif cat:
                        selected.add(p)
                    else:
                        warn(f"Formato '{p.suffix}' não suportado pelo Convertix.")
            else:
                warn(f"Número {idx} fora do intervalo.")


def pick_files() -> list[Path]:
    """Abre direto o navegador visual de arquivos."""
    return browse_files()


def _pick_files_manual() -> list[Path]:
    """Seleção manual por digitação (comportamento original)."""
    import glob as _glob
    step("DIGITAR CAMINHOS")
    info("Digite os caminhos dos arquivos (um por linha).")
    info(f"Pode usar glob como {YE}/pasta/*.jpg{R}")
    info(f"Deixe {YE}vazio{R} e pressione Enter para finalizar.\n")

    files: list[Path] = []
    i = 1
    while True:
        raw = input(f"  {CY}Arquivo {i}{R}: ").strip()
        if not raw:
            if files:
                break
            warn("Informe ao menos 1 arquivo.")
            continue

        matches = _glob.glob(os.path.expanduser(raw))
        if matches:
            for m in matches:
                p = Path(m)
                if p.is_file():
                    files.append(p)
                    ok(f"Adicionado: {p.name}")
                    i += 1
        else:
            p = Path(os.path.expanduser(raw))
            if p.is_file():
                files.append(p)
                ok(f"Adicionado: {p.name}")
                i += 1
            else:
                err(f"Arquivo não encontrado: {raw}")

    return files


def pick_format(files: list[Path]) -> str:
    """Exibe formatos disponíveis e pede escolha."""
    # Descobre quais extensões de saída são possíveis com base nos inputs
    all_outputs = set()
    for f in files:
        ext = f.suffix.lower()
        outs = get_available_outputs(ext)
        all_outputs.update(outs)

    if not all_outputs:
        # Pede ao usuário que digita na mão
        warn("Extensão de entrada desconhecida. Digite a extensão de destino (ex: .mp4):")
        return input(f"  {CY}Formato:{R} ").strip()

    step("ESCOLHER FORMATO DE SAÍDA")

    # Agrupa por categoria
    grouped = {}
    for ext in sorted(all_outputs):
        for cat, info_ in CONVERSION_MAP.items():
            if ext in info_["output"]:
                grouped.setdefault(cat, []).append(ext)
                break

    cat_icons = {
        "image": "🖼  IMAGEM",
        "video": "🎬  VÍDEO",
        "audio": "🎵  ÁUDIO",
        "document": "📄  DOCUMENTO",
        "data": "📊  DADOS",
    }

    options = []
    for cat, exts in grouped.items():
        print(f"\n  {BG_BL}{WH} {cat_icons.get(cat, cat.upper())} {R}")
        for ext in exts:
            n = len(options) + 1
            options.append(ext)
            print(f"    {CY}{n:2d}{R}. {B}{ext}{R}")

    print()
    while True:
        raw = input(f"  {CY}Escolha o número ou digite a extensão:{R} ").strip()
        if raw.isdigit():
            idx = int(raw) - 1
            if 0 <= idx < len(options):
                return options[idx]
            err("Número inválido.")
        elif raw.startswith("."):
            return raw.lower()
        else:
            return "." + raw.lower()


def _render_dir_browser(cwd: Path):
    """Renderiza o navegador de diretórios para escolha de destino."""
    clear()
    print(f"\n{CY}{B}  ⬡ CONVERTIX  {R}{DIM}— Escolher Pasta de Destino{R}")
    print(f"  {DIM}{'─'*62}{R}")

    # Breadcrumb
    parts = list(cwd.parts)
    crumbs = []
    for i, part in enumerate(parts):
        if i == len(parts) - 1:
            crumbs.append(f"{CY}{B} {part} {R}")
        else:
            crumbs.append(f"{DIM}{part}{R}")
    print("  " + f"{DIM} ›{R} ".join(crumbs))
    print(f"  {DIM}{'─'*62}{R}")

    # Atalhos
    print(f"  {DIM}  [n] entrar   [0] subir   [nova] criar pasta   [ok] salvar aqui{R}\n")

    # Lista subpastas
    try:
        subdirs = sorted(
            [p for p in cwd.iterdir() if p.is_dir() and not p.name.startswith(".")],
            key=lambda p: p.name.lower()
        )
    except PermissionError:
        subdirs = []

    if subdirs:
        print(f"  {BL}  📁  SUBPASTAS{R}")
        for i, d in enumerate(subdirs, 1):
            try:
                n = len(list(d.iterdir()))
                sub = f"  {DIM}{n} itens{R}"
            except PermissionError:
                sub = f"  {RE}sem acesso{R}"
            print(f"  {CY}{i:>3}.{R}      {BL}{B}{d.name}/{R}{sub}")
    else:
        print(f"  {DIM}  (sem subpastas){R}")

    print(f"\n  {DIM}{'─'*62}{R}")
    print(f"  {GR}{B}  ► Salvar em:{R}  {WH}{B}{cwd}{R}\n")


def pick_output_dir() -> Path:
    """Navegador interativo para escolher o diretório de saída."""
    cwd = Path.cwd()

    while True:
        # Lista subpastas para navegação
        try:
            subdirs = sorted(
                [p for p in cwd.iterdir() if p.is_dir() and not p.name.startswith(".")],
                key=lambda p: p.name.lower()
            )
        except PermissionError:
            subdirs = []
            cwd = cwd.parent
            continue

        _render_dir_browser(cwd)

        cmd = input(f"  {CY}►{R} ").strip().lower()

        # Confirmar pasta atual
        if cmd in ("ok", "confirmar", "c", ""):
            cwd.mkdir(parents=True, exist_ok=True)
            ok(f"Destino: {cwd}")
            return cwd

        # Subir um nível
        if cmd == "0":
            if cwd.parent != cwd:
                cwd = cwd.parent
            else:
                warn("Já está na raiz.")
            continue

        # Criar nova pasta aqui
        if cmd == "nova" or cmd.startswith("nova "):
            parts = cmd.split(maxsplit=1)
            if len(parts) == 2:
                name = parts[1].strip()
            else:
                name = input(f"  {CY}Nome da nova pasta:{R} ").strip()
            if name:
                new_dir = cwd / name
                try:
                    new_dir.mkdir(parents=True, exist_ok=True)
                    cwd = new_dir
                    ok(f"Pasta criada: {new_dir.name}")
                except Exception as e:
                    err(f"Não foi possível criar: {e}")
                    input(f"  {DIM}Enter para continuar...{R}")
            continue

        # Digitar caminho absoluto
        if cmd.startswith("/") or cmd.startswith("~"):
            p = Path(os.path.expanduser(cmd))
            if p.is_dir():
                cwd = p
            else:
                try:
                    p.mkdir(parents=True, exist_ok=True)
                    cwd = p
                    ok(f"Pasta criada: {p}")
                except Exception as e:
                    err(f"Caminho inválido: {e}")
                    input(f"  {DIM}Enter para continuar...{R}")
            continue

        # Navegar por número
        if cmd.isdigit():
            idx = int(cmd) - 1
            if 0 <= idx < len(subdirs):
                cwd = subdirs[idx]
            else:
                warn(f"Número {cmd} fora do intervalo.")
            continue

        warn("Comando não reconhecido. Use um número, 'ok', '0', 'nova <nome>' ou um caminho.")
        input(f"  {DIM}Enter para continuar...{R}")


def run_conversions(files: list[Path], dst_ext: str, out_dir: Path):
    """Executa todas as conversões com barra de progresso."""
    step("CONVERTENDO ARQUIVOS")
    sep()
    total = len(files)
    results = []

    for i, src in enumerate(files, 1):
        name = src.name
        print(f"\n  {DIM}[{i}/{total}]{R} {WH}{name}{R}  →  {CY}{dst_ext}{R}")
        print(progress_bar(i-1, total))

        success, dst = do_convert(src, dst_ext, out_dir)
        results.append((src, dst, success))

        print(f"\r{progress_bar(i, total)}")
        if success:
            size = dst.stat().st_size if dst.exists() else 0
            ok(f"Salvo: {dst.name}  ({_human(size)})")
        else:
            err(f"Falha ao converter: {name}")

    # Resumo
    sep()
    n_ok  = sum(1 for _,_,s in results if s)
    n_err = total - n_ok
    print(f"\n  {GR}{B}✔ {n_ok} convertido(s){R}   ", end="")
    if n_err:
        print(f"{RE}{B}✘ {n_err} com falha{R}", end="")
    print(f"\n  {DIM}Saída: {out_dir}{R}\n")


def _human(size: int) -> str:
    for unit in ["B","KB","MB","GB"]:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


def show_supported():
    """Lista todos os formatos suportados."""
    step("FORMATOS SUPORTADOS")
    cat_icons = {"image":"🖼","video":"🎬","audio":"🎵","document":"📄","data":"📊"}
    for cat, info_ in CONVERSION_MAP.items():
        icon = cat_icons.get(cat, "•")
        print(f"\n  {icon}  {B}{CY}{cat.upper()}{R}")
        ins  = "  ".join(info_["input"])
        outs = "  ".join(info_["output"])
        print(f"    {DIM}Entrada :{R} {GR}{ins}{R}")
        print(f"    {DIM}Saída   :{R} {YE}{outs}{R}")
    print()


def main_menu():
    """Menu principal."""
    first = True
    while True:
        if not first:
            clear()
        first = False
        banner()
        print(f"""
  {B}O que deseja fazer?{R}

  {CY}1{R}.  {B}Converter arquivos{R}          {DIM}(selecionar 1 ou mais){R}
  {CY}2{R}.  {B}Ver formatos suportados{R}
  {CY}3{R}.  {B}Converter pasta inteira{R}     {DIM}(todos os arquivos de um tipo){R}
  {CY}0{R}.  {DIM}Sair{R}
""")
        choice = input(f"  {CY}►{R} ").strip()

        if choice == "1":
            clear(); banner()
            files = pick_files()
            if not files: continue
            fmt = pick_format(files)
            out = pick_output_dir()
            run_conversions(files, fmt, out)
            input(f"\n  {DIM}Pressione Enter para continuar...{R}")

        elif choice == "2":
            clear(); banner()
            show_supported()
            input(f"  {DIM}Pressione Enter para voltar...{R}")

        elif choice == "3":
            clear(); banner()
            step("CONVERTER PASTA INTEIRA")
            folder = input(f"  {CY}Pasta:{R} ").strip()
            folder = Path(os.path.expanduser(folder))
            if not folder.is_dir():
                err("Pasta não encontrada."); input(); continue

            src_ext = input(f"  {CY}Extensão de origem (ex: .jpg):{R} ").strip()
            if not src_ext.startswith("."): src_ext = "." + src_ext

            files = list(folder.glob(f"*{src_ext}")) + \
                    list(folder.glob(f"*{src_ext.upper()}"))
            if not files:
                warn(f"Nenhum arquivo {src_ext} encontrado em {folder}.")
                input(); continue

            info(f"Encontrado(s): {len(files)} arquivo(s)")
            fmt = pick_format(files)
            out = pick_output_dir()
            run_conversions(files, fmt, out)
            input(f"\n  {DIM}Pressione Enter para continuar...{R}")

        elif choice == "0":
            clear()
            print(f"\n  {CY}Obrigado por usar o {B}Convertix{R}{CY}!{R}\n")
            sys.exit(0)


# ─── CLI modo rápido (argumentos) ─────────────────────────────

def cli_mode():
    """Modo linha de comando: python convertor.py arquivo.mp4 .mp3"""
    import argparse
    parser = argparse.ArgumentParser(
        description="Convertix - Conversor Universal de Arquivos",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  python convertor.py foto.png .jpg
  python convertor.py video.avi .mp4 --output ~/Videos
  python convertor.py *.csv .xlsx
  python convertor.py --list
        """
    )
    parser.add_argument("files", nargs="*", help="Arquivo(s) de entrada")
    parser.add_argument("format", nargs="?", help="Formato de saída (ex: .mp3)")
    parser.add_argument("--output", "-o", help="Diretório de saída")
    parser.add_argument("--list", "-l", action="store_true",
                        help="Lista formatos suportados")

    args = parser.parse_args()

    if args.list:
        banner()
        show_supported()
        return

    if not args.files or not args.format:
        # Sem args: modo interativo
        main_menu()
        return

    banner()
    # Resolve globs
    import glob
    files = []
    for pattern in args.files:
        matches = glob.glob(os.path.expanduser(pattern))
        if matches:
            files.extend(Path(m) for m in matches if Path(m).is_file())
        else:
            p = Path(os.path.expanduser(pattern))
            if p.is_file():
                files.append(p)

    if not files:
        err("Nenhum arquivo encontrado."); sys.exit(1)

    fmt = args.format if args.format.startswith(".") else "." + args.format
    out_dir = Path(os.path.expanduser(args.output)) if args.output \
              else Path.home() / "Convertix_Output"
    out_dir.mkdir(parents=True, exist_ok=True)

    run_conversions(files, fmt, out_dir)


if __name__ == "__main__":
    # Se tiver argumentos além do nome do script → modo CLI
    if len(sys.argv) > 1:
        cli_mode()
    else:
        main_menu()
